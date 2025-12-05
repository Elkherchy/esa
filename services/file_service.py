"""
Service pour la manipulation et l'extraction de contenu des fichiers
"""
import os
import logging
from typing import Optional, Tuple
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument
from PIL import Image

logger = logging.getLogger(__name__)


class FileService:
    """Service pour la gestion et l'extraction de contenu des fichiers"""
    
    # Types de fichiers supportés
    SUPPORTED_DOCUMENT_TYPES = ['.pdf', '.docx', '.doc', '.txt']
    SUPPORTED_IMAGE_TYPES = ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
    
    # Taille maximale de fichier (50 MB)
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    @classmethod
    def validate_file(cls, file) -> Tuple[bool, str]:
        """
        Valide un fichier uploadé
        
        Args:
            file: Fichier Django UploadedFile
            
        Returns:
            Tuple (is_valid, error_message)
        """
        # Vérifier la taille
        if file.size > cls.MAX_FILE_SIZE:
            return False, f"Le fichier est trop volumineux (max {cls.MAX_FILE_SIZE // (1024*1024)} MB)"
        
        # Vérifier l'extension
        file_ext = Path(file.name).suffix.lower()
        all_supported = cls.SUPPORTED_DOCUMENT_TYPES + cls.SUPPORTED_IMAGE_TYPES
        
        if file_ext not in all_supported:
            return False, f"Type de fichier non supporté: {file_ext}"
        
        return True, ""
    
    @classmethod
    def extract_text_from_pdf(cls, file_path: str) -> str:
        """
        Extrait le texte d'un fichier PDF
        
        Args:
            file_path: Chemin vers le fichier PDF
            
        Returns:
            Texte extrait
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction PDF: {e}")
            return ""
    
    @classmethod
    def extract_text_from_docx(cls, file_path: str) -> str:
        """
        Extrait le texte d'un fichier DOCX
        
        Args:
            file_path: Chemin vers le fichier DOCX
            
        Returns:
            Texte extrait
        """
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction DOCX: {e}")
            return ""
    
    @classmethod
    def extract_text_from_txt(cls, file_path: str) -> str:
        """
        Lit le contenu d'un fichier texte
        
        Args:
            file_path: Chemin vers le fichier TXT
            
        Returns:
            Contenu du fichier
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Essayer avec un autre encodage
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Erreur lors de la lecture du fichier texte: {e}")
                return ""
        except Exception as e:
            logger.error(f"Erreur lors de la lecture du fichier: {e}")
            return ""
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """
        Extrait le texte d'un fichier selon son type
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Texte extrait
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return cls.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return cls.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return cls.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Type de fichier non supporté pour l'extraction: {file_ext}")
            return ""
    
    @classmethod
    def get_page_count(cls, file_path: str) -> int:
        """
        Compte le nombre de pages d'un document
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Nombre de pages (0 si non applicable)
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    return len(pdf_reader.pages)
            elif file_ext in ['.docx', '.doc']:
                # Pour DOCX, on estime environ 500 mots par page
                doc = DocxDocument(file_path)
                total_words = sum(len(p.text.split()) for p in doc.paragraphs)
                return max(1, total_words // 500)
            else:
                return 1  # Fichiers texte = 1 page
        except Exception as e:
            logger.error(f"Erreur lors du comptage des pages: {e}")
            return 0
    
    @classmethod
    def generate_snippet(cls, text: str, max_length: int = 200) -> str:
        """
        Génère un extrait d'un texte
        
        Args:
            text: Texte source
            max_length: Longueur maximale de l'extrait
            
        Returns:
            Extrait du texte
        """
        if not text:
            return ""
        
        # Nettoyer le texte
        text = " ".join(text.split())
        
        if len(text) <= max_length:
            return text
        
        # Couper au dernier espace avant max_length
        snippet = text[:max_length]
        last_space = snippet.rfind(' ')
        
        if last_space > 0:
            snippet = snippet[:last_space]
        
        return snippet + "..."
    
    @classmethod
    def get_file_info(cls, file_path: str) -> dict:
        """
        Récupère les informations d'un fichier
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Dictionnaire avec les informations du fichier
        """
        try:
            file_stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            return {
                'size': file_stat.st_size,
                'extension': file_ext,
                'page_count': cls.get_page_count(file_path),
            }
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des infos du fichier: {e}")
            return {
                'size': 0,
                'extension': '',
                'page_count': 0,
            }


